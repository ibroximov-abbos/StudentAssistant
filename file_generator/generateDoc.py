from docx import Document
from docx.shared import Pt,Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

def create_titul(full_name, subject, theme):
    organisation = "O’ZBEKISTON RESPUBLIKASI OLIY TA’LIM VA INOVATSIYALAR VAZIRLIGI MUHAMMAD AL-XORAZMIY NOMIDAGI TOSHKENT AXBOROT TEXNALOGIYALARI UNIVERSITETI"
    paragraph = doc.add_paragraph(organisation)
    run = paragraph.runs[0]
    font = run.font
    font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    font.size = Pt(16)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.5

    doc.add_paragraph()

    try:
        doc.add_picture('file_generator/logo.png', width=Inches(3), height=Inches(3))  
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print("Rasm yuklanmadi:", e)
    paragraph = doc.add_paragraph('\n\n')
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(theme.capitalize())
    run.font.name = 'Times New Roman'
    run.bold = True
    run.font.size = Pt(16)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph = doc.add_paragraph('\n\n\n\n')
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"Fan: {subject}\nTalaba: {full_name}")
    font = run.font
    font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    font.size = Pt(14)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def create_doc(data, full_name, username, subject, theme):
    create_titul(full_name, subject, theme)  

    doc.add_page_break()

    paragraph = doc.add_paragraph(data)
    run = paragraph.runs[0]
    font = run.font
    font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    font.size = Pt(14)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5

    filename = f"{username}.docx"
    doc.save(filename)
    print(f"Fayl saqlandi: {filename}")
    return 'ok'
